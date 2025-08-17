#!/usr/bin/env python3
"""
Enhanced document processing module for Kitobxon Kids Bot
Multi-format document import and export with optimized performance
"""

import asyncio
import logging
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

# Optional document processing imports
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer, PageBreak
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import PyPDF2
    PDF_TEXT_AVAILABLE = True
except ImportError:
    try:
        import pdfplumber
        PDF_TEXT_AVAILABLE = True
    except ImportError:
        PDF_TEXT_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Enhanced document processor with multi-format support"""
    
    def __init__(self):
        self.supported_formats = {
            'text/plain': ['.txt'],
            'application/pdf': ['.pdf'],
            'application/msword': ['.doc'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx'],
            'application/vnd.ms-excel': ['.xls'],
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': ['.xlsx']
        }
    
    def validate_document(self, filename: str, mime_type: str, file_size: int) -> Tuple[bool, str]:
        """Validate document format and size"""
        max_size = 50 * 1024 * 1024  # 50MB
        
        if file_size > max_size:
            return False, f"Fayl hajmi juda katta ({file_size / 1024 / 1024:.1f} MB). Maksimal hajm: 50 MB"
        
        if not filename:
            return False, "Fayl nomi ko'rsatilmagan"
        
        file_ext = Path(filename).suffix.lower()
        
        # Check by extension
        for mime, extensions in self.supported_formats.items():
            if file_ext in extensions:
                return True, "OK"
        
        # Check by MIME type
        if mime_type in self.supported_formats:
            return True, "OK"
        
        return False, f"Qo'llab-quvvatlanmaydigan format: {file_ext or mime_type}"
    
    async def extract_text_from_document(self, file_content: bytes, filename: str, mime_type: str) -> Tuple[bool, str]:
        """Extract text from various document formats"""
        try:
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.txt' or mime_type == 'text/plain':
                return await self._extract_from_text(file_content)
            elif file_ext in ['.docx'] and DOCX_AVAILABLE:
                return await self._extract_from_docx(file_content)
            elif file_ext in ['.xlsx', '.xls'] and EXCEL_AVAILABLE:
                return await self._extract_from_excel(file_content)
            elif file_ext == '.pdf' and PDF_TEXT_AVAILABLE:
                return await self._extract_from_pdf(file_content)
            else:
                return False, f"Matn ajratish qo'llab-quvvatlanmaydi: {file_ext}"
                
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            return False, f"Hujjatdan matn ajratishda xatolik: {str(e)}"
    
    async def _extract_from_text(self, file_content: bytes) -> Tuple[bool, str]:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'cp1251', 'ascii']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return True, text.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with errors='ignore'
            text = file_content.decode('utf-8', errors='ignore')
            return True, text.strip()
            
        except Exception as e:
            return False, f"Matn faylini o'qishda xatolik: {str(e)}"
    
    async def _extract_from_docx(self, file_content: bytes) -> Tuple[bool, str]:
        """Extract text from Word document"""
        try:
            doc = DocxDocument(BytesIO(file_content))
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_parts)
            return True, full_text
            
        except Exception as e:
            return False, f"Word hujjatini o'qishda xatolik: {str(e)}"
    
    async def _extract_from_excel(self, file_content: bytes) -> Tuple[bool, str]:
        """Extract text from Excel file"""
        try:
            workbook = openpyxl.load_workbook(BytesIO(file_content), data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = [f"=== {sheet_name} ==="]
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None and str(cell_value).strip():
                            row_text.append(str(cell_value).strip())
                    
                    if row_text:
                        sheet_text.append(' | '.join(row_text))
                
                if len(sheet_text) > 1:  # More than just the header
                    text_parts.extend(sheet_text)
                    text_parts.append('')  # Empty line between sheets
            
            full_text = '\n'.join(text_parts)
            return True, full_text
            
        except Exception as e:
            return False, f"Excel faylini o'qishda xatolik: {str(e)}"
    
    async def _extract_from_pdf(self, file_content: bytes) -> Tuple[bool, str]:
        """Extract text from PDF file"""
        try:
            # Try pdfplumber first if available
            try:
                import pdfplumber
                pdf_file = BytesIO(file_content)
                text_parts = []
                
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text.strip())
                
                full_text = '\n\n'.join(text_parts)
                return True, full_text
                
            except ImportError:
                # Fallback to PyPDF2
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                text_parts = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text.strip())
                
                full_text = '\n\n'.join(text_parts)
                return True, full_text
                
        except Exception as e:
            return False, f"PDF faylini o'qishda xatolik: {str(e)}"
    
    def parse_test_questions(self, text: str) -> Tuple[bool, List[Dict[str, Any]]]:
        """Parse test questions from extracted text"""
        try:
            questions = []
            
            # Try different question patterns
            patterns = [
                r'(\d+)\.?\s*(.+?)\n\s*[aA]\)\s*(.+?)\n\s*[bB]\)\s*(.+?)\n\s*[cC]\)\s*(.+?)\n\s*[dD]\)\s*(.+?)\n\s*(?:Javob|Answer|Ответ)[:=]?\s*([aAbBcCdD])',
                r'(\d+)\.?\s*(.+?)\n.*?[aA][\).]?\s*(.+?)\n.*?[bB][\).]?\s*(.+?)\n.*?[cC][\).]?\s*(.+?)\n.*?[dD][\).]?\s*(.+?)\n.*?(?:Javob|Answer|Ответ)[:=]?\s*([aAbBcCdD])',
                r'Savol\s*(\d+)[:.]?\s*(.+?)\n.*?A[:.]?\s*(.+?)\n.*?B[:.]?\s*(.+?)\n.*?C[:.]?\s*(.+?)\n.*?D[:.]?\s*(.+?)\n.*?(?:Javob|Answer|Ответ)[:=]?\s*([aAbBcCdD])'
            ]
            
            for pattern in patterns:
                matches = re.findall(pattern, text, re.MULTILINE | re.DOTALL)
                
                if matches:
                    for match in matches:
                        question_num, question_text, option_a, option_b, option_c, option_d, correct_answer = match
                        
                        # Clean up text
                        question_text = self._clean_text(question_text)
                        options = [
                            self._clean_text(option_a),
                            self._clean_text(option_b),
                            self._clean_text(option_c),
                            self._clean_text(option_d)
                        ]
                        
                        # Convert answer to index
                        answer_map = {'a': 0, 'b': 1, 'c': 2, 'd': 3}
                        correct_index = answer_map.get(correct_answer.lower(), 0)
                        
                        if question_text and all(options):
                            questions.append({
                                'question': question_text,
                                'options': options,
                                'correct_answer': correct_index
                            })
                    break
            
            if not questions:
                # Try simple Q&A format
                qa_pattern = r'(?:Savol|Question|Вопрос)\s*(\d+)[:.]?\s*(.+?)\n(?:Javob|Answer|Ответ)[:.]?\s*(.+?)(?=\n(?:Savol|Question|Вопрос)|\n\n|\Z)'
                qa_matches = re.findall(qa_pattern, text, re.MULTILINE | re.DOTALL)
                
                for match in qa_matches:
                    question_num, question_text, answer_text = match
                    
                    question_text = self._clean_text(question_text)
                    answer_text = self._clean_text(answer_text)
                    
                    if question_text and answer_text:
                        # Create multiple choice from the answer
                        options = [answer_text, f"Variant B", f"Variant C", f"Variant D"]
                        questions.append({
                            'question': question_text,
                            'options': options,
                            'correct_answer': 0
                        })
            
            return True, questions
            
        except Exception as e:
            logger.error(f"Error parsing questions: {str(e)}")
            return False, []
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        
        # Remove extra whitespace and newlines
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Remove common formatting artifacts
        text = re.sub(r'[^\w\s\-.,!?()[\]{};"\'`~@#$%^&*+=|\\/<>]', '', text)
        
        return text.strip()
    
    async def export_to_excel(self, data: List[Dict[str, Any]], filename: str = "export.xlsx") -> Tuple[bool, bytes]:
        """Export data to Excel format with professional formatting"""
        if not EXCEL_AVAILABLE:
            return False, b"Excel export qo'llab-quvvatlanmaydi"
        
        try:
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "Ma'lumotlar"
            
            if not data:
                return False, b"Export qilish uchun ma'lumot yo'q"
            
            # Set up styles
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center")
            
            # Border style
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Write headers
            headers = list(data[0].keys())
            for col, header in enumerate(headers, 1):
                cell = sheet.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = thin_border
            
            # Write data
            for row_idx, record in enumerate(data, 2):
                for col_idx, (key, value) in enumerate(record.items(), 1):
                    cell = sheet.cell(row=row_idx, column=col_idx, value=str(value))
                    cell.border = thin_border
                    
                    # Auto-adjust column width
                    column_letter = get_column_letter(col_idx)
                    if len(str(value)) > 15:
                        sheet.column_dimensions[column_letter].width = min(50, len(str(value)) + 2)
                    else:
                        sheet.column_dimensions[column_letter].width = max(15, len(str(value)) + 2)
            
            # Save to bytes
            excel_buffer = BytesIO()
            workbook.save(excel_buffer)
            excel_buffer.seek(0)
            
            return True, excel_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error exporting to Excel: {str(e)}")
            return False, b"Excel export xatoligi"
    
    async def export_to_pdf(self, data: List[Dict[str, Any]], title: str = "Ma'lumotlar", filename: str = "export.pdf") -> Tuple[bool, bytes]:
        """Export data to PDF format with professional formatting"""
        if not PDF_AVAILABLE:
            return False, b"PDF export qo'llab-quvvatlanmaydi"
        
        try:
            pdf_buffer = BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
            
            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=18,
                spaceAfter=30,
                alignment=1  # Center
            )
            
            story = []
            
            # Title
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))
            
            if not data:
                story.append(Paragraph("Ma'lumot topilmadi", styles['Normal']))
            else:
                # Prepare table data
                headers = list(data[0].keys())
                table_data = [headers]
                
                for record in data:
                    row = []
                    for key in headers:
                        value = str(record.get(key, ""))
                        # Truncate long text for better formatting
                        if len(value) > 50:
                            value = value[:47] + "..."
                        row.append(value)
                    table_data.append(row)
                
                # Create table
                table = Table(table_data)
                
                # Table style
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(table)
            
            # Add timestamp
            story.append(Spacer(1, 30))
            timestamp = datetime.now().strftime("%d.%m.%Y %H:%M")
            story.append(Paragraph(f"Yaratilgan: {timestamp}", styles['Normal']))
            
            doc.build(story)
            pdf_buffer.seek(0)
            
            return True, pdf_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error exporting to PDF: {str(e)}")
            return False, b"PDF export xatoligi"
    
    async def export_test_results_excel(self, results: List[Dict[str, Any]]) -> Tuple[bool, bytes]:
        """Export test results to Excel with specialized formatting"""
        if not EXCEL_AVAILABLE:
            return False, b"Excel export qo'llab-quvvatlanmaydi"
        
        try:
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "Test Natijalari"
            
            # Headers
            headers = [
                "T/r", "Bola ismi", "Ota-ona ismi", "Viloyat", "Tuman", 
                "Mahalla", "Yoshi", "Telefon", "Test sanasi", "Ball", 
                "Foiz", "Baho", "Holati"
            ]
            
            # Set up styles
            header_font = Font(bold=True, color="FFFFFF", size=12)
            header_fill = PatternFill(start_color="2F5597", end_color="2F5597", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            
            # Status colors
            pass_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
            fail_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
            
            # Write headers
            for col, header in enumerate(headers, 1):
                cell = sheet.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
            
            # Write data
            for idx, result in enumerate(results, 2):
                sheet.cell(row=idx, column=1, value=idx-1)
                sheet.cell(row=idx, column=2, value=result.get('child_name', ''))
                sheet.cell(row=idx, column=3, value=result.get('parent_name', ''))
                sheet.cell(row=idx, column=4, value=result.get('region', ''))
                sheet.cell(row=idx, column=5, value=result.get('district', ''))
                sheet.cell(row=idx, column=6, value=result.get('mahalla', ''))
                sheet.cell(row=idx, column=7, value=result.get('age', ''))
                sheet.cell(row=idx, column=8, value=result.get('phone', ''))
                sheet.cell(row=idx, column=9, value=result.get('test_date', ''))
                sheet.cell(row=idx, column=10, value=result.get('score', 0))
                
                percentage = result.get('percentage', 0)
                sheet.cell(row=idx, column=11, value=f"{percentage:.1f}%")
                sheet.cell(row=idx, column=12, value=result.get('grade', ''))
                
                status_cell = sheet.cell(row=idx, column=13, value="O'tdi" if result.get('passed', False) else "O'tmadi")
                if result.get('passed', False):
                    status_cell.fill = pass_fill
                else:
                    status_cell.fill = fail_fill
            
            # Auto-adjust column widths
            for column in sheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                
                adjusted_width = min(50, max(12, max_length + 2))
                sheet.column_dimensions[column_letter].width = adjusted_width
            
            # Save to bytes
            excel_buffer = BytesIO()
            workbook.save(excel_buffer)
            excel_buffer.seek(0)
            
            return True, excel_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error exporting test results to Excel: {str(e)}")
            return False, b"Excel export xatoligi"

# Global document processor instance
document_processor = DocumentProcessor()
